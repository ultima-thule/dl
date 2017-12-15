from mongoengine import *
import lib.mongoLeankit

import argparse
import datetime
import lib.confluence

import docx
from docx.shared import Inches
from lxml import html
from lxml.html.clean import clean_html
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
import sys
import time

import credentials
import lib.confluence
import lib.confluence_parser
import io

gl_font_size = 11
gl_font_name = "Calibri"

def _initMongoConn ():
    connect('leankit')

def add_simple_par(document, text, size=11, align=docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT, bold=False, breakpar=False):
    if document is not None:
        global gl_font_name

        paragraph = document.add_paragraph()
        paragraph.paragraph_format.page_break_before = breakpar
        paragraph.alignment = align

        run = paragraph.add_run()
        run.font.bold = bold
        run.font.name = gl_font_name
        run.font.size = docx.shared.Pt(size)
        run.add_text(text)


def add_simple_cell(table, row, cell, text, size=11, align=docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT, bold=False):
    if table is not None:
        cell = table.cell(row, cell)
        cell.text = text


def generate_pw_first_page(document, project_name):
    if document is not None:

        font_size_title = 16

        for i in range(0, 8):
            add_simple_par(document, "", font_size_title)

        add_simple_par(document, "Załącznik nr 1", font_size_title, align=docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER)
        add_simple_par(document, "do", font_size_title, align=docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER)
        add_simple_par(document, "Porozumienia Wykonawczego nr", font_size_title, align=docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER)
        add_simple_par(document, "\"" + project_name + "\"", font_size_title, align=docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER, bold=True)

        for i in range(0, 8):
            add_simple_par(document, "", font_size_title)

        table = document.add_table(rows=3, cols=2)
        table.autofit = True
        table.alignment = WD_TABLE_ALIGNMENT.RIGHT

        font_size_table = 10

        add_simple_cell(table, 0, 0, "Autor", font_size_table)
        add_simple_cell(table, 0, 1, "", font_size_table)
        add_simple_cell(table, 1, 0, "Data utworzenia", font_size_table)
        add_simple_cell(table, 1, 1, datetime.datetime.now().strftime("%Y-%m-%d"), font_size_table)
        add_simple_cell(table, 2, 0, "Wersja", font_size_table)
        add_simple_cell(table, 2, 1, "1.0", font_size_table)

        generate_pw_toc(document)


def generate_pw_toc(document):
    global gl_font_size
    add_simple_par(document, "", gl_font_size, align=docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER, bold=False, breakpar=True)
    document.add_heading("Spis treści", 1)

    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')  # creates a new element
    fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
    instrText.text = 'TOC \\o "1-2" \\h \\z \\u'  # change 1-3 depending on heading levels you need

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = "Kliknij prawym klawiszem myszki, aby zaktualizować spis treści."
    fldChar2.append(fldChar3)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)
    p_element = paragraph._p


def generate_pw_last_page(document):
    if document is not None:
        global gl_font_size

        document.add_heading("Produkty prac", 1)
        txt = "Realizacja zadań sprintu, wymagań projektu – wszystkie elementy, które definiują przyrost funkcjonalności, "
        txt += "niezbędny z punktu widzenia wartości biznesowej produktu, definiowanej przez Product Ownera. "
        txt += "W ramach prac przewidzianych niniejszą ofertą przewidziano dostarczenie m. in. "
        txt += "produktów prac opisanych w rozdziale Kamienie milowe."
        add_simple_par(document, txt, gl_font_size, align=docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY)

        document.add_heading("Uwagi dot. harmonogramu i kosztorysu", 1)
        txt = "Zlecający biorąc aktywny udział w wycenach poszczególnych zadań w projekcie "
        txt += "zobowiązuje się do akceptacji lub odrzucenia wyceny zadań zgodnie ze współczynnikiem wydajności "
        txt += "określanym podczas ww. spotkań, najdalej w następnym dniu "
        txt += "po spotkaniu. Prace postępować będą w trybie dwutygodniowych sprintów."
        add_simple_par(document, txt, gl_font_size, align=docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY)

        txt = "Termin zakończenia projektu uzależniony jest od długości trwania testów przeprowadzanych przez "
        txt += "Zamawiającego i w przypadku konieczności wydłużenia tego czasu, termin zamknięcia projektu przesuwa "
        txt += "się odpowiednio, jednak nie dłużej niż 10 dni roboczych od planowanej daty wdrożenia."
        add_simple_par(document, txt, gl_font_size, align=docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY)


def generate_page_project_card(document, page_tree):
    if document is not None and page_tree is not None:
        global gl_font_size
        add_simple_par(document, "", gl_font_size, align=docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER, bold=False, breakpar=True)
        document.add_heading("Podstawowe informacje o projekcie", 1)

        panels = page_tree.xpath("//acmacro/acparameter[@acname='title']")
        for p in panels:
            # Clean title from CCF formatting
            heading = p.text.replace("*", "")
            if heading == "Checklista" or heading== "Uczestnicy projektu" or heading == "Kosztorys" \
                    or heading == "Dokumentacja projektowa":
                return
            document.add_heading(heading, 2)

            # Generate text
            text_body = p.xpath("./following-sibling::acrich-text-body/*")
            for tb in text_body:
                generate_paragraph(document, tb, 0, None, None, None)

            # Generate tables
            tables = p.xpath("./following-sibling::acrich-text-body/table")
            table = generate_table(document, tables, table_format=True)


def generate_page_sprint(document, page_tree, heading):
    if document is not None and page_tree is not None:
        document.add_heading(heading, 2)
        tables = page_tree.xpath("//div/table")
        generate_table(document, tables, table_format=False)


def generate_page_close(document, page_tree):
    if document is not None and page_tree is not None:
        document.add_heading("Zakończenie projektu", 1)

    panels = page_tree.xpath("//acmacro/acparameter[@acname='title']")
    for p in panels:
        # Generate text
        text_body = p.xpath("./following-sibling::acrich-text-body/*")
        for tb in text_body:
            generate_paragraph(document, tb, 0, None, None, None)

        # Generate tables
        tables = p.xpath("./following-sibling::acrich-text-body/table")
        table = generate_table(document, tables, table_format=True)


def generate_paragraph(document, item, li_indent=0, prev_par=None, current_par=None, cell=None, bold=False, addFormatting=False):
    # do not parse tables - separate function!
    if item.tag == "table" or item.tag == "acparameter" or item.tag == "acplaceholder":
        return

    global gl_font_size
    global gl_font_name

    indent = li_indent

    if prev_par is None:
        if cell is not None:
            prev_par = cell.add_paragraph()
        else:
            prev_par = document.add_paragraph()
    if current_par is None:
        current_par = prev_par
    current_run = None

    # opening tags and text
    if item.tag is not None:
        current_run = current_par.add_run()
        current_run.font.name = gl_font_name
        current_run.font.size = docx.shared.Pt(gl_font_size)

        tag_lowered = item.tag.lower()

        if tag_lowered == "b" or tag_lowered == "strong" or bold is True:
            current_run.font.bold = True
        elif tag_lowered == "i":
            current_run.font.italic = True
        elif tag_lowered == "s":
            current_run.font.strike = True
        elif tag_lowered == "li":
            prev_par = current_par
            if cell is not None:
                current_par = cell.add_paragraph(style='List Bullet')
            else:
                current_par = document.add_paragraph(style='List Bullet')
            if indent != 0:
                current_par.paragraph_format.left_indent = Inches(0.5)
            current_run = current_par.add_run()
            indent += 1
        elif tag_lowered == "br":
            prev_par = current_par
            if cell is not None:
                current_par = cell.add_paragraph()
            else:
                current_par = document.add_paragraph()
            current_run = current_par.add_run()
        elif tag_lowered == "ripage":
            current_run.add_text(item.attrib["ricontent-title"])

    txt = item.text
    if txt is not None and txt != "":
        if addFormatting is True:
            txt += ": "
        current_run.add_text(txt)

    # all child elements
    for elem in item:
        generate_paragraph(document, elem, indent, prev_par, current_par, cell, bold, addFormatting)

    # closing tags
    if item.tag is not None:
        current_run = current_par.add_run()
        current_run.font.name = gl_font_name
        current_run.font.size = docx.shared.Pt(gl_font_size)

        tag_lowered = item.tag.lower()

        if tag_lowered == "b" or tag_lowered == "strong" or bold is True:
            current_run.font.bold = None
        elif tag_lowered == "i":
            current_run.font.italic = None
        elif tag_lowered == "s":
            current_run.font.strike = None
        elif tag_lowered == "li":
            current_par = prev_par
            current_run = current_par.add_run()

    # tail text
    if item.tail is not None:
        current_run.add_text(item.tail)

    return prev_par


def generate_table(document, tables, table_format=True):
    if tables is not None and document is not None:
        for t in tables:
            h1 = t.xpath("../preceding-sibling::h1[1]")
            doc_table = []
            ths = t.xpath("./tbody/tr/th")
            if len(ths) == 0:
                ths = t.xpath("./thead/tr/th")
            if len(ths) == 0:
                continue

            # do nothing if table contains sprint details or cost info
            if len(ths) > 0 and ths[0].text == "Kod sprintu":
                continue
            if len(ths) > 1 and ths[1].text == "Suma SP na sprint":
                continue

            for th in ths:
                doc_table.append([th])
            tds = t.xpath("./tbody/tr/td")
            not_empty = len(tds)==0
            for index, elem in enumerate(tds):
                i = index % len(ths)
                doc_table[i].append(elem)
                not_empty |= (elem.text is not None and elem.text.strip() != "" and elem.text.strip() != "n/a") or elem.tail is not None

            # do nothing if table contains nothing besides the headers
            if not_empty is False:
                continue

            # add section header
            if len(h1) > 0:
                document.add_heading(h1[0].text_content(), 3)

            convert_table_to_word(document, doc_table, table_format)
            add_simple_par(document, "")


def convert_table_to_word(document, parsed_table, table_format=True):
    # convert to word
    cols = len(parsed_table)
    if cols > 0:
        rows = len(parsed_table[0])
        if table_format is True:
            table = document.add_table(rows=rows, cols=cols)
            table.autofit = True
            table.style = 'Light List Accent 1'

            for i, col in enumerate(parsed_table):
                for j, elem in enumerate(col):
                    generate_paragraph(document, elem, 0, None, None, table.cell(j, i))
        else:
            cols_no = len(parsed_table)
            rows_no = len(parsed_table[0])
            skip_column = set()
            prev_par =  None

            for r in range(0, rows_no):
                # skip specific tasks
                elem_title = parsed_table[0][r]
                txt = elem_title.text_content().lower()
                if "spotkania scrumowe" in txt or "zarządzanie projektem" in txt or "grooming" in txt:
                    continue

                for c in range(0, cols_no):
                    if len(parsed_table[c]) > r:
                        elem = parsed_table[c][r]

                        # skip unnecessary columns
                        txt = elem.text_content().lower()
                        if "story point" in txt or "czas realizacji" in txt or "czy zrealizowane?" in txt:
                            skip_column.add(c)

                        # do not print first (header) row and skipped columns
                        if r > 0 and c not in skip_column:
                            prev_par = generate_paragraph(document, parsed_table[c][0], 0, None, None, None, bold=True, addFormatting=True)
                            generate_paragraph(document, elem, 1, prev_par, None, None, bold=False)

                # empty line between rows
                add_simple_par(document, "")


def check_for_project_card (cf_server, parent_page):
    # iterate through children
    child_pages = cf_server.get_child_pages(parent_page)
    for page in child_pages:
        title_doc = page["title"].partition("|")[0]
        title = page["title"].lower()

        if "karta projektu" in title or "kp" in title:
            return True

    return False


def get_page_tree(cf_server, page):
    page_full = cf_server.get_page_by_id(page["id"])
    page_tree = ""
    if page_full is not None:
        page_clean = page_full["content"].replace("ac:", "ac").replace("ri:", "ri")
        page_tree = html.fromstring(page_clean)

    return page_tree


if __name__ == '__main__':
    if len(sys.argv) < 2:
        exit("Usage: " + sys.argv[0] + " projectname")
    project_name = sys.argv[1]

    _initMongoConn()

    #init Confluence connection
    user_cf = credentials.loginConfluence['consumer_secret']
    pwd_cf = credentials.loginConfluence['password']
    url = "http://doc.grupa.onet/rpc/xmlrpc"
    confl = lib.confluence.Confluence(url, user_cf, pwd_cf, "PROJEKTY")

    date_text = datetime.datetime.now().strftime("%Y%m%d-%H%M%S")
    parent_page_title = project_name

    document = docx.Document()

    project_page = confl.get_page(project_name)
    #print(project_page)
    if project_page is not None:
        #create title PW pages
        generate_pw_first_page(document, project_name)

        #create project card from main page
        project_card_exists = check_for_project_card(confl, project_page)
        if not project_card_exists:
            page_tree = get_page_tree(confl, project_page)
            generate_page_project_card(document, page_tree)

        # get child pages and sort them
        child_pages = confl.get_child_pages(project_page)
        pages_structure = {"sprint_cards": []}
        for page in child_pages:
            title_doc = page["title"].partition("|")[0]
            title = page["title"].lower()

            page_tree = get_page_tree(confl, page)

            if "sprint" in title:
                pages_structure["sprint_cards"].append(page_tree)
            elif "zakończenie projektu" in title:
                pages_structure["project_close"] = page_tree
            elif "karta projektu" in title or "kp" in title:
                pages_structure["project_card"] = page_tree

        # iterate pages and generate doc
        # project card
        page_content = pages_structure.get("project_card", None)
        generate_page_project_card(document, page_content)

        # sprint cards
        add_simple_par(document, "", gl_font_size, align=docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False, breakpar=True)
        document.add_heading("Zakres realizowanych prac", 1)
        txt = "W ramach projektu zostaną wykonane zadania, opisane poniżej w ramach kolejnych sprintów."
        add_simple_par(document, txt, gl_font_size, align=docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY)

        sprint_pages = pages_structure.get("sprint_cards")
        sprint_no = 1
        for page_content in sprint_pages:
            generate_page_sprint(document, page_content, "Sprint # " + str(sprint_no))
            sprint_no += 1

        # closure card
        #page_content = pages_structure.get("project_close", None)
        #generate_page_close(document, page_content)

        #create last PW pages
        generate_pw_last_page(document)

    # workaround na przekazywanie pliku do frontu
    document.save("./pw_files/%s.docx" % project_name)

    # wylaczam, robie workaround jak powyżej
    #target_stream = io.BytesIO()
    #document.save(target_stream)
    #target_stream.seek(0)
    #pwfile = lib.mongoLeankit.Pwfile()

    #pwfile.data = target_stream.read()
    #pwfile.project = project_name
    #pwfile.generation_date = datetime.datetime.now()
    #pwfile.date_text = date_text
    #pwfile.format_type = "DOCX"
    #pwfile.save()
    #target_stream.close()
    exit(0)
