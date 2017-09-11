from html.parser import HTMLParser
from docx import Document

class ConffHtmlParser(HTMLParser):

    def __init__(self, document):
        HTMLParser.__init__(self)
        self.document = document

    def handle_starttag(self, tag, attrs):
        # self.document.add_paragraph("Encountered a start tag:", tag)
        print("Encountered a start tag:", tag)

    def handle_endtag(self, tag):
        # self.document.add_paragraph("Encountered an end tag:", tag)
        print("Encountered an end tag :", tag)

    def handle_data(self, data):
        # self.document.add_paragraph("Encountered some data  :", data)
        print("Encountered some data  :", data)