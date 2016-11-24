import argparse
import datetime
import lib.confluence

from docx import Document
from lxml import html
from lxml.html.clean import clean_html

import credentials
import lib.confluence
import lib.confluence_parser

def generate_page_project_card(document, page):
    if document is not None and page is not None:
        page_clean = page["content"].replace("ac:", "ac").replace("ri:", "ri")
        tree = html.fromstring(page_clean)
        panels = tree.xpath("//acmacro/acparameter[@acname='title']")
        for p in panels:
            heading = p.text.replace("*", "")
            document.add_heading(heading)
            text_body = p.xpath("./following-sibling::acrich-text-body/*")
            for tb in text_body:
                paragraphs = tb.xpath("./text()")
                generate_paragraph(document, paragraphs)
            tables = p.xpath("./following-sibling::acrich-text-body/table")
            generate_table(document, tables)

def generate_page_sprint(document, page):
    if document is not None and page is not None:
        page_clean = page["content"].replace("ac:", "ac").replace("ri:", "ri")
        tree = html.fromstring(page_clean)
        tables = tree.xpath("//div/table")
        generate_table(document, tables)

def generate_page_close(document, page):
    pass


def generate_paragraph(document, paragraphs):
    if paragraphs is not None and document is not None:
        for par in paragraphs:
            document.add_paragraph(par)

def generate_table(document, tables):
    if tables is not None and document is not None:
        for t in tables:
            h1 = t.xpath("../preceding-sibling::h1[1]")
            if len(h1) > 0:
                document.add_heading(h1[0].text_content())
            doc_table = []
            ths = t.xpath("./tbody/tr/th")
            if len(ths) == 0:
                ths = t.xpath("./thead/tr/th")
            for th in ths:
                doc_table.append([th.text_content()])
            tds = t.xpath("./tbody/tr/td")
            for index, elem in enumerate(tds):
                i = index % len(ths)
                doc_table[i].append(elem.text_content())

            # convert to word
            cols = len(doc_table)
            if cols > 0:
                rows = len(doc_table[0])
                table = document.add_table(rows=rows, cols=cols)
                table.style = 'Light List Accent 1'

                for i, col in enumerate(doc_table):
                    for j, row in enumerate(col):
                        table.cell(j,i).text = row

                document.add_paragraph()


if __name__ == '__main__':

    # parse arguments
    parser = argparse.ArgumentParser(description='Generate xlsx estimate sheeto for project.')
    parser.add_argument('-p', '--project', help='project name', required=True)
    memory_parser = parser.add_mutually_exclusive_group(required=True)
    memory_parser.add_argument('--memory', dest='memory', action='store_true')
    memory_parser.add_argument('--no-memory', dest='memory', action='store_false')
    parser.set_defaults(memory=True)
    args = parser.parse_args()

    #init Confluence connection
    user_cf = credentials.loginConfluence['consumer_secret']
    pwd_cf = credentials.loginConfluence['password']
    url = "http://doc.grupa.onet/rpc/xmlrpc"
    confl = lib.confluence.Confluence(url, user_cf, pwd_cf, "PROJEKTY")

    file_name = args.project + "_zakres_" + datetime.datetime.now().strftime("%Y%m%d-%H%M%S.docx")

    parent_page_title = args.project

    document = Document()

    project_page = confl.get_page(args.project)
    if project_page is not None:
        child_pages = confl.get_child_pages(project_page)
        for page in child_pages:
            title = page["title"].lower()
            page_full = confl.get_page_by_id(page["id"])

            if "karta projektu" in title:
                generate_page_project_card(document, page_full)
            elif "sprint" in title:
                generate_page_sprint(document, page_full)
            elif "zakończenie projektu" in title:
                generate_page_close(document, page_full)

    document.save(file_name)

    # if args.memory:
    #     connect('leankit')
    #     report = lib.mongoLeankit.Estimate()
    #
    #     report.xls_data = data
    #     report.generation_date = datetime.datetime.now()
    #     report.project_name = args.project
    #
    #     report.save()

    exit(0)
