import argparse
import datetime
import lib.confluence

import docx
from docx.shared import Inches
from lxml import html
from lxml.html.clean import clean_html
import sys

import credentials
import lib.confluence
import lib.confluence_parser
from mongoengine import *
import lib.mongoLeankit
import io

def _initMongoConn ():
    connect('leankit')


def generate_page_project_card(document, page_tree):
    if document is not None and page_tree is not None:
        panels = page_tree.xpath("//acmacro/acparameter[@acname='title']")
        for p in panels:
            # Clean title from CCF formatting
            heading = p.text.replace("*", "")
            document.add_heading(heading, 2)

            # Generate text
            text_body = p.xpath("./following-sibling::acrich-text-body/*")
            for tb in text_body:
                generate_paragraph(document, tb, 0, None, None, None)

            # Generate tables
            tables = p.xpath("./following-sibling::acrich-text-body/table")
            generate_table(document, tables)


def generate_page_sprint(document, page_tree):
    if document is not None and page_tree is not None:
        tables = page_tree.xpath("//div/table")
        generate_table(document, tables)


def generate_page_close(document, page_tree):
    pass


def generate_paragraph(document, item, li_indent=0, prev_par=None, current_par=None, cell=None):
    # do not parse tables - separate function!
    if item.tag == "table" or item.tag == "acparameter" or item.tag == "acmacro":
        return

    indent = li_indent

    if prev_par is None:
        if cell is not None:
            prev_par = cell.add_paragraph()
        else:
            prev_par = document.add_paragraph()
    if current_par is None:
        current_par = prev_par
    current_run = None

    # opening tags and text
    if item.tag is not None:
        current_run = current_par.add_run()
        if item.tag.lower() == "b" or item.tag.lower() == "strong":
            current_run.font.bold = True
        elif item.tag.lower() == "i":
            current_run.font.italic = True
        elif item.tag.lower() == "s":
            current_run.font.strike = True
        elif item.tag.lower() == "li":
            prev_par = current_par
            if cell is not None:
                current_par = cell.add_paragraph(style='List Bullet')
            else:
                current_par = document.add_paragraph(style='List Bullet')
            if indent != 0:
                current_par.paragraph_format.left_indent = Inches(0.5)
            current_run = current_par.add_run()
            indent += 1
    if item.text is not None and item.text != "":
        current_run.add_text(item.text)

    # all child elements
    for elem in item:
        generate_paragraph(document, elem, indent, prev_par, current_par, cell)

    # closing tags
    if item.tag is not None:
        current_run = current_par.add_run()
        if item.tag.lower() == "b" or item.tag.lower() == "strong":
            current_run.font.bold = None
        elif item.tag.lower() == "i":
            current_run.font.italic = None
        elif item.tag.lower() == "s":
            current_run.font.strike = None
        elif item.tag.lower() == "li":
            current_par = prev_par
            current_run = current_par.add_run()

    # tail text
    if item.tail is not None:
        current_run.add_text(item.tail)

    return prev_par


def generate_table(document, tables):
    if tables is not None and document is not None:
        for t in tables:
            h1 = t.xpath("../preceding-sibling::h1[1]")
            if len(h1) > 0:
                document.add_heading(h1[0].text_content(), 2)
            doc_table = []
            ths = t.xpath("./tbody/tr/th")
            if len(ths) == 0:
                ths = t.xpath("./thead/tr/th")
            for th in ths:
                doc_table.append([th])
            tds = t.xpath("./tbody/tr/td")
            for index, elem in enumerate(tds):
                i = index % len(ths)
                doc_table[i].append(elem)

            # convert to word
            cols = len(doc_table)
            if cols > 0:
                rows = len(doc_table[0])
                table = document.add_table(rows=rows, cols=cols)
                table.autofit = True
                table.style = 'Light List Accent 1'

                for i, col in enumerate(doc_table):
                    for j, elem in enumerate(col):
                        generate_paragraph(document, elem, 0, None, None, table.cell(j,i))
                document.add_paragraph()

def check_for_project_card (cf_server, parent_page):
    # iterate through children
    child_pages = cf_server.get_child_pages(parent_page)
    for page in child_pages:
        title_doc = page["title"].partition("|")[0]
        title = page["title"].lower()

        if "karta projektu" in title or "kp" in title:
            return True

    return False


def get_page_tree(cf_server, page):
    page_full = cf_server.get_page_by_id(page["id"])
    page_clean = page_full["content"].replace("ac:", "ac").replace("ri:", "ri")
    page_tree = html.fromstring(page_clean)

    return page_tree


if __name__ == '__main__':
    if len(sys.argv) < 2:
        exit("Usage: " + sys.argv[0] + " projectname")
    project_name = sys.argv[1]

    _initMongoConn()

    #init Confluence connection
    user_cf = credentials.loginConfluence['consumer_secret']
    pwd_cf = credentials.loginConfluence['password']
    url = "http://doc.grupa.onet/rpc/xmlrpc"
    confl = lib.confluence.Confluence(url, user_cf, pwd_cf, "PROJEKTY")

    date_text = datetime.datetime.now().strftime("%Y%m%d-%H%M%S")

    parent_page_title = project_name

    document = docx.Document()

    project_page = confl.get_page(project_name)
    if project_page is not None:
        project_card_exists = check_for_project_card(confl, project_page)
        #create project card from main page
        if not project_card_exists:
            page_tree = get_page_tree(confl, project_page)
            document.add_heading(project_page["title"], 1)
            generate_page_project_card(document, page_tree)

        #iterate through children
        child_pages = confl.get_child_pages(project_page)
        for page in child_pages:
            title_doc = page["title"].partition("|")[0]
            title = page["title"].lower()

            page_tree = get_page_tree(confl, page)

            if "karta projektu" in title or "kp" in title:
                document.add_heading(title_doc, 1)
                generate_page_project_card(document, page_tree)
            elif "sprint" in title:
                document.add_heading(title_doc, 1)
                generate_page_sprint(document, page_tree)
            elif "zakończenie projektu" in title:
                document.add_heading(title_doc, 1)
                generate_page_close(document, page_tree)


    target_stream = io.BytesIO()
    document.save(target_stream)
    target_stream.seek(0)
    pwfile = lib.mongoLeankit.Pwfile()

    pwfile.data = target_stream.read()
    pwfile.project = project_name
    pwfile.generation_date = datetime.datetime.now()
    pwfile.date_text = date_text
    pwfile.format_type = "DOCX"
    pwfile.save()

    target_stream.close()

    exit(0)
